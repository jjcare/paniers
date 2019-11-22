# -*- coding: utf-8 -*-
################################################################################
#
#  PANIERS DE NOËL (pydocxfamilles.py, formerly paniers2012.py)
#
#       John Carey, 2010, 2011, 2012, 2015, 2018
#
#  pydocx version : complet rewrite (2018) to take data directly from
#                   MS-Word document tables. No need to add terminator to
#                   original data.
#
#                   - find best (most recent) .docx candidate in current
#                     directory (.doc file must be resaved as .docx)
#
#                   - read in table information and read into data structures
#
#                   - geolocate all addresses for mapping purposes
# 
#                   - produce output .sql and .tab files of data
#
#                   REQUIRES: foyers-épiceries.txt, foyers-familles.txt
#
#                   TODO: some data massaging for typos (misspelled streets)
#                         more error checking
#                         make table creation sql
#                         join all sql files into one (less importing)
#
#
#################################################################################

# info comes in tables from word doc (must be converted to docx)

from docx import Document
import os, re, sys, unicodedata, json, time, datetime
from urllib import request

# global variables
fname = 'familles.txt'
outfiles = 'famrecs.tab famrecs.sql deprecs.tab deprecs.sql'.split()
three_mos = 60*60*24*28*3  # 3 months in seconds
debug = 0  # print debugging information (0 to go faster
testdata = 0   # when 1, skip location look-up (good for testing the input data)

# amounts for calculating certificates (global values presently)
base_amt = 60      # monoparental amount
couple_deduct = 10 # reduce by this if couple
depend_amt = 10    # amt for each dependent

def strip_accents(s): # remove accents for url requests (ascii only)
   return u''.join(c for c in unicodedata.normalize('NFD', s)
                  if unicodedata.category(c) != 'Mn') 

def getFamilyList():
    # find the latest file to open
    poss = sorted([ x for x in os.listdir() if ".docx" in x],
                  key=lambda x:os.stat(x).st_ctime)
    if poss:
        print('Liste de familles trouvée : {}\n'.format(poss[-1]))
        return poss[-1] # latest candidate
    return '' # no candidate


class Error(Exception):
    pass

class DataError(Error):
        def __init__(self,message):
            self.message = message

class WrongNumberError(Error):
        def __init__(self,message):
            self.message = message

class HTTPError(Error):
        def __init__(self,message):
            self.message = message


#    
#  class structures
#


class Foyers: # foyer assignments

    def __init__(self):
        self.foyers = {}
        # chick if last year's list is still there
        if time.time() - os.stat("foyers-familles.txt").st_mtime > three_mos:  # data is old
            print (">>>> Fichier 'foyers-familles.txt' est vieux. Foyers ne sont pas inclus. Utilisez makeFoyersFamilles().")
            return
        try:
            print ( "Élaboration des foyers en cours...")
            f = open('foyers-familles.txt','r')
            buf = f.read()
            buf = re.split('\r?\n', buf)
            if buf[0].find('Foyer') != 0:
                print ("Foyer header not found")
                return  # foyer should be first
            buf.pop(0)                            # don't count field names
            # use dict comprehension to populate foyer list (strip preceding 0s)
            self.foyers = {i.split('\t')[0]: list(map(lambda x: x.lstrip('0'),re.split('[^0-9]+',i.split('\t')[1]))) for i in buf if i.strip() != ''}
        except:
            print ("Foyer list failed")
            return

    def addFoyer (self, s):
        if len(s) == 1:
            self.foyers[s[0]] = ""
        else:
            self.foyers[s[0]] = re.split('[^0-9]+', s[1]) # accept any sep char

    def getFoyer (self, fam):
        # search for family number fam in foyer list. If more than one family for foyer, get index (eg. 203a)
        for foyer in self.foyers:
            if fam[0] == '0': fam = fam[1:]
            if fam in self.foyers[foyer] or '0'+fam in self.foyers[foyer]:
                if len(self.foyers[foyer]) == 1:
                    app = ''
                else:
                    app = 'abcde'[self.foyers[foyer].index(fam)]
                return foyer + app
        return ''

    def getAllFoyers (self):
        return sorted([ x for x in self.foyers ])

    def getFamily (self, foyer):
        if self.foyers.has_key(foyer):
            return foyer + ' : ' + ' + '.join(self.foyers[foyer])
        return ''
# end class foyer

class Geo:
    """Keeps track of geolocation values for addresses"""

    geohash = {}
    hashfname = "addresses.txt"

    def __init__(self):
        # only read in hash table once
        if Geo.geohash: return
        # checks for hashed list of addresses and loads it
        if not os.path.exists(Geo.hashfname):  # create it
            f = open(Geo.hashfname, 'w', 1, 'utf-8')
            json.dump({},f)
            f.close()
        hp = open(Geo.hashfname, "r")
        # deserialize dict from hashes file  
        Geo.geohash = json.load(hp, parse_int=int)
        hp.close()

    def add(self, key, val):
        key = str(key)  # make sure key is string
        if not key in Geo.geohash:
            Geo.geohash[key] = val

    def __get__(self, instance, key):
        return Geo.geohash.get(key)
 
    def geolocate (self, addr):
        # google geolocate api request
        if testdata:
            return '00,00'
        if debug: print ('geolocate: ', addr)
        addr = strip_accents(addr)  # normalize to ascii
        location = self.getLatLong(addr)
        if not location:
            # try without postcode
            location = self.getLatLong ( re.split(' [A-Z][0-9][A-Z] [0-9][A-Z][0-9]', addr)[0])
            if not location:
                # try just postcode
                location = self.getLatLong ( addr[-7:] )
                if not location:
                    print ('Warning: no result for geocode (%s)' % (addr,))
                    location = ''
        print ('.', end=' ')
        return location            

    def getLatLong (self, addr):
        # get google geolocate data
        addr = '+'.join(addr.split(' ')) # spaces to '+'
        if debug: print ('\nSearch location = "{}"'.format(addr))

        # check in address location table first
        if addr in Geo.geohash:
            return Geo.geohash[addr]
        else:
            print('+', end ='')
            try:
                location = self.google2loc(addr)
            except HTTPError:
                print('Geolocate server error. Retry job later.')
                return False

            self.add( addr, location)
            time.sleep(.3)  # wait
            return location
        
    def google2loc (self, address):
        google_api_key = 'AIzaSyCqJgDfAWRHX4iGg4Jbd3KU0uScQQps_lA'
        google_url = 'https://maps.googleapis.com/maps/api/geocode/json?address={},+Montreal,+Quebec,+Canada&key={}'
        response = request.urlopen(google_url.format(address.replace(' ','+'), google_api_key))
        j = json.loads(response.read().decode())
        if len(j) and 'error' not in j:
           try:
              lat = j['results'][0]['geometry']['location']['lat']
              lon = j['results'][0]['geometry']['location']['lng']
              return '{},{}'.format(lat,lon)
           except:
              print('\nAdresse pas trouvé : ',address)
        return ''

    def writeGeo(self):
        # write out the updated hash table of address locations
        f = open(Geo.hashfname, 'w', encoding='utf-8')
        json.dump(Geo.geohash, f)
        f.close()

    def __delete__(self):
        self.writeGeo()

# end class Geo


# under IDLE, no need. os.chdir('/media/john/Lexar/paniers de noel')

class FamilyBuffer:
    """FamilyBuffer : get list of families and process into list of FamilyRecord,
    doing error correction, formatting and exposing inspection methods."""

    def __init__(self, docname):
        from docx import Document
        if not docname: raise InputError('FamilyBuffer::No family list document given.')
        doc = Document(docname)

        # make list of table rows (ignore headers and empty rows)
        fams = []
        for i in range(len(doc.tables)):
            t = doc.tables[i]
            for j in range(len(t.rows)):
                r = [ c.text for c in t.row_cells(j) ]
                if re.match(r'[0-9]+',r[0]):  # don't keep header lines
                    fams.append(r)

        # compensate for bug in doc to docx conversion (table cells are repeated)
        #  -- remove duplicate cell items

        for i,f in enumerate(fams):
            fams[i] = [v for j,v in enumerate(fams[i]) if j==0 or v != fams[i][j-1]]


        # sanity checks

        if len(fams) == 0: # uh-oh - what, no data? Stop the press!
            print("Aucune famille trouvée. Vérifie la source de données.")
            raise DataError('Aucune famille trouvée')

        if len(set([len(f) for f in fams])) != 1: # all should have the same field count
            raise DataError('Longueurs de données inégales')
        
        rep = input(u"J'ai trouvé {} familles. Est-ce juste (o/n)? ".format(len(fams)))
        if rep[0].lower() == 'n':
            raise WrongNumberError('Vérifiez le format des données.')  # crash out if data looks suspicous

        self.fams = fams  # keep the result if good

    def display(self,fam):
        print('\nFamille no. {}\n'.format(fam[0]), file=sys.stderr)
        for i in fam[1:]: print('{} ##'.format(i))

    def showFams(self):
        print ("showFam - '.' to end, or fam number.\n")
        for i,f in enumerate(self.fams):
            x = input('? ')
            if x == '.': break
            if x.isdigit() and int(x) < len(self.fams):
                self.display(self.fams[int(x)-1])
            else: self.display(f)

class FamilyRecords:

    # make a list of family records
    def __init__(self, fambuf, foyers):
        self.famlist = []
        g = Geo() # get saved list of geocodes

        for f in fambuf.fams:
            self.famlist.append(self.FamilyRecord(f, foyers))

        g.writeGeo() # save list of geocodes

    class FamilyRecord:

        # class constant definitions
        UPPER_NAME_PAT = r"[A-ZÉÈÙÂÄÀÇÎÏÔÖ -][']?[A-ZÉÈÙÂÄÀÇÎÏÔÖ -]{2,}\s" # corrected for false results (eg. ST-LOUIS = ST, LOUIS and N'Doube = N', Doube)
        SEX_PAT = r'^([MF])$'
        ADD_PAT = r'^([0-9]+? [^0-9]+( *# ?[0-9]+)?)'
        ADDR_PAT = r'^([0-9]{2,5} ?(?:[A-F] )?) ?([^0-9#]+) *#? ?([0-9]+)?'  # grab all components of address at once
        PCODE_PAT = r'^H\d{1}[A-Z]{1} *\d{1}[A-Z]{1}\d{1}$'  # for MTL. For generic: r'^[ABCEGHJKLMNPRSTVXY]{1}\d{1}[A-Z]{1} *\d{1}[A-Z]{1}\d{1}$'
        TEL_PAT = r'^(\?|\*?[-0-9]{7,})'
        UPPER_PAT = r"[A-ZÉÈÙÂÄÀÇÎÏÔÖ -][']?[A-ZÉÈÙÂÄÀÇÎÏÔÖ -]{2,}\s"

        def __init__(self, fam, foyers):
            fields = ['num','members','nip','addr','pcode','tels','children', 'sexes', 'ages']
            
            funcs = {'members':self.setMembers,
                     'addr':self.setAddress,
                     'tels':self.setTels,
                     'children':self.setChildren,
                     'sexes': self.setSexes,
                     'ages':self.setAges }
            
            self.montant = 0
            
            # check for record length (should have been previously verified)
            if len(fam) and len(fam) < len(fields):
                raise DataError('Famille {}: record trop court.'.format(fam[0]))

            
            # iterate all input fields and set attributes accordingly
            for i,f in enumerate(fam):
                if fields[i] in funcs:
                    f = funcs[fields[i]](f) # call appropriate transformation function
                setattr(self, fields[i],f)

            # other information
            self.fixChildren() # combine children data into dict form.
            self.setMontant()
            self.ra = ''
            self.foyer = foyers.getFoyer(self.num)
            geo = Geo()  # get a Geo instance 
            self.location = geo.geolocate(' '.join([self.addr['number'],self.addr['rue'],self.pcode]))

        def setMembers(self, fam):
            # sets the Demandeur and other adults (children are handled elsewhere)
            members = []
            fam = fam.split('\n')
            name, sex = self.getNextMember(fam) # process name and sex
            
            while name:
                
                if len(members) == 0:   # demandeur is always first name
                    rel = 'Demandeur'
                elif re.search('\d{2}', name):  # age given in name indicates other person
                    rel = 'Autre'
                    if not (20 < int(re.search('\d{2}', name).group(0)) < 60): self.montant += depend_amt
                else:
                    rel = 'Conjoint'
                    self.montant -= couple_deduct
                    
                nom, prenom = self.getNameParts(name)
                members.append({ 'nom': nom, 'prenom': prenom, 'sexe': sex, 'relation':rel })
                name, sex = self.getNextMember(fam)
            return members
    
        def getNextMember(self, fam):
            # returns a tuple of name, sex
            if fam:
                mem = fam.pop(0).strip()
                if '..' in mem: # get rid of ...... between name and sex
                   mem = ''.join(mem.split('.'))

                if mem and mem[-1] in 'MF':
                    sex = mem[-1]
                    mem = mem[:-1]
                else:
                    sex = ''
                
                if re.match(self.UPPER_NAME_PAT, mem,re.IGNORECASE):
                    return (mem.strip(), sex)
            return ('', '')

        def getNameParts(self, name):
            w = re.match(self.UPPER_PAT, name)
            if not w:
                if ' ' in name:
                    nom, prenom = name.split(' ',1) # consider first word as name
                    nom = nom.upper()   # todo: make this international (locale dependent)
                else:
                    nom = name
                    prenom = ''
            else:
                nom = w.group(0).strip()
                prenom = name[w.end():]
            return (nom, prenom)

        def setAddress(self, address):
            # parses and sets the street address
            address = address.split('\n')[0] # ignore any second lines
            
            if not address[:3].strip().isdigit(): raise ValueError ("Address expected, not found: {}".format(address))

            w = re.match(self.ADDR_PAT, address.strip())
            if w:
                adparts = w.groups()
                addr = {}
                addr['number'] = adparts[0].strip()
                rue = adparts[1].strip()
                addr['rue'] = rue if rue[-1] != '-' else rue[:-1].strip()
                if 'app.' in addr['rue'][-4:]: addr['rue'] = addr['rue'][:-4].strip()
                addr['appt'] = adparts[2].strip() if adparts[2] else 'null'
            else:
                raise ValueError ('Invalid address: %s' % (address,))

            return addr
        
        def setTels(self,tels):
            # sets the telephone numbers
            telnos = []
            tels = tels.split('\n')
            
            while tels and re.match(self.TEL_PAT, tels[0].strip()):
                telnos.append(self.formatTelno(tels.pop(0)))

            if len(telnos) == 0 : raise SystemError ("setTels: No telephone numbers found in family #%s: %s" % (self.num, str(tels)))
            return telnos

        def formatTelno(self, t):
            if t[0] == '*': t = '438' + t[1:]
            t = re.sub('[^0-9]', '', t)
            if len(t) < 10: t = '514' + t
            return format(int(t[:-1]), ",").replace(",", "-") + t[-1]  # not mine - can't think of better just now


        def setChildren(self,ch):
            # parse and set children

            children = []
            ch = ch.split('\n')

            cnames = [ {'prenom':n.strip()} for n in ch if n.strip]

            return cnames
            

        def setSexes(self,sexes):
            sexes = [s.strip() for s in sexes.split('\n') if s.strip() in 'MF']
            if len(sexes) != len(self.children):
                raise IndexError('setSexes: incorrect number of items')
            return sexes
            

        def setAges (self, ages):
            ages = [format_age(age) for age in ages.split('\n') if age.strip()]
            if len(ages) != len(self.children):
                raise IndexError ('setAges: not enough ages.')
            return ages

        def fixChildren(self):
            # join children data into dict form
            for i in range(len(self.children)):
                self.children[i]['sexe'] = self.sexes[i]
                self.children[i]['age'] = self.ages[i]
                self.children[i]['relation'] = 'Enfant'
            del self.sexes
            del self.ages


        def setFoyer(self, num):
            # set foyer
            self.foyer = num


        def setMontant(self):
            # based on family membership, calculate amount to be given
            self.montant += base_amt + (depend_amt * len(self.children))

        # other methods to display, get and transform family records


        def showFamily (self):
            sys.stderr.write( '\n' + self.num + '  ---------------------\n') # colour heading
            for i in 'members nip showAddress() pcode location tel ra showChildren() foyer montant'.split():
                print ('\t' + i.capitalize() + ':', eval('self.' + i))

        def showAddress (self):
            ad = self.addr
            pd = '#' if ad['appt'] != 'null' else ''
            return '{} {} {}{}'.format(ad['number'], ad['rue'], pd, ad['appt'])

        def showChildren (self):
            ret = '\n'
            for child in self.children:
                ret = ret + ('\t\t{}\t{}  {} ans\n'.format (child['prenom'],child['sexe'],child['age']))
            return ret
            
        def getFamilyRec(self):
            addr = '%(number)s\t%(rue)s\t%(appt)s' % self.addr
            tel = self.tels[0] + '\t' + ','.join(self.tels[1:])
            return '%s\t%s\t%s\t%s\t%s\t%s\t%d\t%s\t%s\n' % (self.num, self.nip, addr, self.pcode, tel, self.ra, self.montant, self.foyer, self.location)

        def getFamilySql(self):
            addr = '"%(number)s","%(rue)s","%(appt)s"' % self.addr
            if len(self.tels) == 1:
                self.tels.append("")
            tel = self.tels[0] + '","' + ','.join(self.tels[1:])
            return '(%s,%s,%s,"%s","%s","%s","%d", "%s", "%s"),\n' % (self.num, self.nip, addr, self.pcode, tel, self.ra, self.montant, self.foyer, self.location)
            
        def getDependentsRec(self):
            rec = ''
            for i in self.members:
                rec = rec + self.nip + '\t'
                rec = rec + '%(relation)s\t%(nom)s\t%(prenom)s\t\t%(sexe)s\n' % i
            for i in self.children:
                rec = rec + self.nip + '\t'
                rec = rec + '%(relation)s\t\t\t%(prenom)s\t%(age)s\t%(sexe)s\n' % i
            return rec

        def getDependentsSql(self):
            rec = '('
            for i in self.members:
                rec = rec + self.nip + ','
                rec = rec + '"%(relation)s","%(nom)s","%(prenom)s",null,"%(sexe)s"),\n(' % i
            for i in self.children:
                rec = rec + self.nip + ','
                rec = rec + '"%(relation)s",null,"%(prenom)s","%(age)s","%(sexe)s"),\n(' % i
            return rec[:-1]

    # end class FamilyRecord

def doGroceries():  # get the grocery list produce sql commands and text files
    if time.time() - os.stat("foyers-épiceries.txt").st_mtime <= three_mos:  # data is recent
        groceries = getGroceries("foyers-épiceries.txt")  # dict of foyers - items
        # make sql commands
        outbuf = "use shalom_cnd;\nSET character_set_client = utf8;\nTRUNCATE TABLE groceries;\nINSERT INTO groceries (`foyer`, `item1`, `item2`) VALUES \n"
        outbuf += '\n'.join(['({},"{}","{}"),'.format(foyer,groceries[foyer].get('item1'),groceries[foyer].get('item2')) for foyer in groceries])
        outbuf = outbuf[:-1] + ';'
        with open("groceries.sql","w", encoding='utf8') as f:
           f.write(outbuf)

        # make groceries.txt (foyer, item1, item2)
        out = '\n'.join(['\t'.join([x, groceries[x].get('item1'), groceries[x].get('item2')])
                         for x in groceries])
        with open("groceries.txt","w", encoding='utf8') as f:
           f.write(out)

        # make grocery-view-parents.txt for letter to parents
        rev = {}
        for foyer in groceries:
           key = '\t'.join([groceries[foyer]['item1'], groceries[foyer].get('item2','')])
           if not rev.get(key):
              rev[key] = [foyer]
           else:
              rev[key].append(foyer)
        out = '\n'.join(['{}\t{}'.format(', '.join(rev[x]), x) for x in rev])
        with open('grocery-view-parents.txt','w',encoding='utf8') as f:
           f.write(out)

        
def getGroceries(fname): # return a dictionary of groceries by foyer
   with open(fname, encoding='utf8') as fh:
      recs = [x for x in fh.read().split('\n') if x.strip()]
    
   foyers={}

   def pluriel (s):
       words = { 'boite':'boites', 'sac':'sacs', 'emballage':'emballages',
                 'pot':'pots', 'bouteille':'bouteilles'}
       return ' '.join([words.get(x,x) for x in s.split(' ')])

   for r in recs:
       rec = r.split('\t')
       rec[3] = pluriel(rec[3]) if int(rec[4]) > 1 else rec[3]
       item = '{} ({} {})'.format(rec[2],rec[4], rec[3])
       if foyers.get(rec[0]): # already there
           foyers[rec[0]]['item2'] = item
       else:
           foyers[rec[0]] = {'item1':item, 'item2':''}
   return foyers
 
def makeFoyerFamilles():
   # use this function to produce foyers-familles.txt
   g = getGroceries("foyers-épiceries.txt")
   i = 1
   fams = []
   niv = input('Premier niveau pour 2 familles : ')
   if int(niv) not in range(6): return
   niv = int(niv) - 1
   for foyer in sorted(g.keys()):
      if int(foyer[0]) > niv:
         fams.append('{}\t{}+{}'.format(foyer, i, i+1))
         i += 2
      else:
         fams.append('{}\t{}'.format(foyer, i))
         i += 1
   with open('foyers-familles.txt', 'w', encoding='utf8') as fh:
      fh.write('\n'.join(fams))
   print('{} familles distribuées.'.format(fams[-1].split('+')[-1]))
         
# for age conversion
conv = {'¾': 0.75, '½': 0.5, '¼': 0.25, '⅓': 0.33}    

def format_age(a):
    # format_age: convert ½, ¼, ¾, ⅓, etc to month values
    # return age else convert to month values if fractions present (unless over 2yrs)
    if '/' in a:
        age = int(a.split('/')[0]) / int(a.split('/')[1] * 12)
        s = 's' if age > 1 else ''
        return '{} mois'.format(age,s)

    if a[-1] not in conv: return a
    comp = list(a)
    frac = round(conv.get(comp.pop()) * 12)
    units = '0' + ''.join(comp)
    if int(units) > 2:
        return str(int(units) + frac)
    return '{} mois'.format( round((int(units) * 12) + frac))    

def toc():
    return datetime.datetime.utcnow()

def elapsed(td):
    m,s = divmod(td.seconds,60)
    return '%d minutes %d seconds' % (m,s)

#####################
#
#  MAIN
#
#####################

t_start = toc()

doGroceries()   # get the grocery list (in 'foyers-épicerie.txt' from FMPro

# read in Family list (docx format) and create FamilyBuffer
try:
    fambuf = FamilyBuffer(getFamilyList())

except DataError(message):
    print("DataError: {}.".format(message))

except WrongNumberError:
    print("Nombre incorrect de familles. Essaie encore.", file=sys.stderr)

except:
    print("Erreur non gérée.")
    raise
else:
    print("Données reçues avec succès. Élaboration des données")
          
    foyers = Foyers()              # get foyer list - if present

    # format data into records, doing error checking and corrections
    #  - MOST OF THE WORK DONE HERE
    families = FamilyRecords(fambuf, foyers)

    # 
    # output results in various formats
    #

    out = [ '','','','']
    total = 0

    # sql famrecs.sql  deprecs.sql
    insertcmd = {'famrecs.sql':'USE shalom_cnd;\nSET character_set_client = utf8;\nTRUNCATE TABLE families;\nINSERT INTO families (`famno`, `nip`, `number`, `street`, `appt`, `code`, `tel_h`, `tel_alt`, `note`, `montant`, `foyer`, `location`) VALUES ',
                 'deprecs.sql':'USE shalom_cnd;\nSET character_set_client = utf8;\nTRUNCATE TABLE dependents;\nINSERT INTO dependents (`nip`, `relation`, `nom`, `prenom`, `age`, `sexe`) VALUES ' }


    for f in families.famlist:
        out[0] += f.getFamilyRec()
        out[1] += f.getFamilySql()
        out[2] += f.getDependentsRec()
        out[3] += f.getDependentsSql()
        total += f.montant

    for i in range(len(out)):
        f = open(outfiles[i], 'w', 1, 'utf-8')
        if 'sql' in outfiles[i]:
            out[i] = insertcmd[outfiles[i]] + out[i][:-2] + ';'
        f.write(out[i])
        f.close()

    # create all-purpose mysql file for project (shalom_cnd.sql + all outputs combined)
    try:
       with open('shalom_cnd.sql', encoding='utf8') as f:
          outbuf = f.read() + '\n'
       if time.time() - os.stat("groceries.sql").st_mtime <= three_mos:
          with open('groceries.sql', encoding='utf8') as f:
             outbuf += f.read() + '\n'
       outbuf += '\n'.join([out[1],out[3]])

       with open('paniers.sql', 'w', encoding='utf8') as f:
          f.write(outbuf)
       print ('\n\nComplete database constructor in « paniers.sql ».\n\n')
    except:
       print ('\n\n==> Unable to make database constructor.')

    print ('\n\nDone. Outputs are in', ', '.join(outfiles), '.\n')
    print ('Total needed for this year\'s collection: %0.2d $\n' % (total,))
    print ('Time:', time.strftime('%X - %x', time.localtime()))
    # get the elapsed time
    t_stop = toc()
    print (elapsed(t_stop - t_start))

finally:
    print('Au revoir.')


